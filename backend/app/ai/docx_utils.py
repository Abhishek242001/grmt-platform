"""Shared .docx-opening helper — used by both grammar_check.py and
format_compliance_check.py, since both need to open .docx files and both
can hit the same real-world compatibility gap (planning log §33).

Some real .docx files (confirmed against a genuine third-party-generated
template, not authored by Microsoft Word directly) use an old/alternate
"purl.oclc.org" transitional OOXML namespace instead of the canonical
"schemas.openxmlformats.org" one that python-docx's Package loader requires
exactly. Without this fallback, such files fail to open at all — not a
formatting-quality issue, a hard crash on otherwise-valid, readable content.
"""
import io
import re
import zipfile

from docx import Document
from docx.document import Document as DocumentObject

# purl.oclc.org/ooxml/{category}/{rest} -> schemas.openxmlformats.org/{category}/2006/{rest}
# Verified against a real file's actual namespace declarations (relationships,
# wordprocessingml, drawingml, math all follow this exact pattern).
_LEGACY_OOXML_NS = re.compile(rb'http://purl\.oclc\.org/ooxml/([^/"]+)/')


def _normalize_legacy_ooxml_namespace(file_path: str) -> io.BytesIO:
    """Rewrites every XML part inside the .docx zip to use the canonical
    OOXML namespace, entirely in memory — no temp file written to disk."""
    output = io.BytesIO()
    with zipfile.ZipFile(file_path, "r") as zin, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith((".xml", ".rels")):
                data = _LEGACY_OOXML_NS.sub(rb"http://schemas.openxmlformats.org/\1/2006/", data)
            zout.writestr(item, data)
    output.seek(0)
    return output


def open_docx(file_path: str) -> DocumentObject:
    """Opens a .docx, falling back to namespace normalization only if the
    standard open fails — genuine Word-generated files (the overwhelming
    majority) pay no extra cost; only the rare non-standard file does."""
    try:
        return Document(file_path)
    except KeyError:
        normalized = _normalize_legacy_ooxml_namespace(file_path)
        return Document(normalized)


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_TXBX_CONTENT_TAG = f"{_W_NS}txbxContent"
_PARAGRAPH_TAG = f"{_W_NS}p"
_TEXT_TAG = f"{_W_NS}t"


def extract_textbox_paragraphs(doc: DocumentObject) -> list[str]:
    """Returns the text of every paragraph living inside a text box —
    content `doc.paragraphs` never sees at all, since it's nested inside a
    `w:txbxContent` element rather than being a direct child of the document
    body (confirmed empirically: `paragraph.text`/`run.text` only walk a
    run's DIRECT `w:t` children, never descending into a nested shape).

    This isn't a rare edge case for real IEEE-formatted submissions: the
    IEEE conference template's OWN official guidance explicitly recommends
    inserting figures via a text box ("more stable than directly inserting
    a picture"), so a figure/table caption placed exactly the way IEEE
    tells authors to place it is otherwise silently invisible to every
    check that reads document text. Confirmed against a real submission
    (table/figure check calibration, Aug 2026) — the check flagged "Figure
    1 referenced but no caption found" and the same for Table 1, and in
    both cases the caption genuinely existed in the file, just inside a
    text box grammar_check.py's extraction never looked at.

    Covers both the legacy VML wrapper (`<v:textbox>`) and the modern
    DrawingML wrapper — both nest their content in a `w:txbxContent`
    element under the same wordprocessingml namespace, so one XML walk for
    that tag name covers both without branching on text-box flavor."""
    paragraphs = []
    for txbx in doc.element.body.iter(_TXBX_CONTENT_TAG):
        for p in txbx.iter(_PARAGRAPH_TAG):
            text = "".join(t.text or "" for t in p.iter(_TEXT_TAG))
            if text.strip():
                paragraphs.append(text)
    return paragraphs
