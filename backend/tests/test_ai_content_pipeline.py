from unittest.mock import patch

import pytest
from docx import Document

from app.ai.ai_content_pipeline import aggregate_chunk_results, run_pipeline


# ── aggregate_chunk_results — pure logic, hand-verifiable ───────────

def test_percentage_is_word_weighted_not_a_flat_average():
    # Independently hand-verified: 10-word flagged chunk out of 210 total
    # words = 4.76% — the whole point of word-weighting is that a SMALL
    # flagged chunk barely moves the percentage, unlike a flat average
    # where it would count exactly as much as a 100-word chunk.
    chunks = [
        {"word_count": 10, "ai_probability": 0.9},
        {"word_count": 100, "ai_probability": 0.1},
        {"word_count": 100, "ai_probability": 0.1},
    ]
    result = aggregate_chunk_results(chunks, chunk_probability_threshold=0.5, max_ai_percentage=15.0)
    assert result["ai_word_count"] == 10
    assert result["total_word_count"] == 210
    assert result["ai_generated_percentage"] == pytest.approx(4.761904761904762)
    assert result["overall_verdict"] == "accept"  # 4.76% < 15%


def test_large_flagged_chunk_moves_percentage_proportionally_more():
    # Same shape as above but the FLAGGED chunk is the large one this time —
    # confirms word-weighting responds to actual chunk size, not just
    # "how many chunks are flagged" (both cases flag exactly 1 of 3 chunks).
    chunks = [
        {"word_count": 100, "ai_probability": 0.9},
        {"word_count": 100, "ai_probability": 0.1},
        {"word_count": 100, "ai_probability": 0.1},
    ]
    result = aggregate_chunk_results(chunks, chunk_probability_threshold=0.5, max_ai_percentage=15.0)
    assert result["ai_generated_percentage"] == pytest.approx(33.333333333333336)
    assert result["overall_verdict"] == "reject"  # 33.3% >= 15%


def test_flags_only_chunks_above_probability_threshold():
    chunks = [
        {"word_count": 50, "ai_probability": 0.6},
        {"word_count": 50, "ai_probability": 0.3},
        {"word_count": 50, "ai_probability": 0.9},
    ]
    result = aggregate_chunk_results(chunks, chunk_probability_threshold=0.5, max_ai_percentage=15.0)
    assert result["flagged_chunk_indices"] == [0, 2]
    assert result["flagged_chunk_count"] == 2
    assert result["ai_word_count"] == 100


def test_accept_requires_strictly_below_max_percentage():
    # "Must have less than 15%" — exactly AT 15% must fail, not pass. Built
    # to land on exactly 15.0% by construction (15 AI words / 100 total).
    chunks = [
        {"word_count": 15, "ai_probability": 0.9},
        {"word_count": 85, "ai_probability": 0.1},
    ]
    result = aggregate_chunk_results(chunks, chunk_probability_threshold=0.5, max_ai_percentage=15.0)
    assert result["ai_generated_percentage"] == pytest.approx(15.0)
    assert result["overall_verdict"] == "reject"


def test_accept_just_under_max_percentage():
    chunks = [
        {"word_count": 14, "ai_probability": 0.9},
        {"word_count": 86, "ai_probability": 0.1},
    ]
    result = aggregate_chunk_results(chunks, chunk_probability_threshold=0.5, max_ai_percentage=15.0)
    assert result["ai_generated_percentage"] == pytest.approx(14.0)
    assert result["overall_verdict"] == "accept"


def test_zero_percent_when_nothing_flagged():
    chunks = [
        {"word_count": 100, "ai_probability": 0.1},
        {"word_count": 100, "ai_probability": 0.2},
    ]
    result = aggregate_chunk_results(chunks, chunk_probability_threshold=0.5, max_ai_percentage=15.0)
    assert result["ai_generated_percentage"] == pytest.approx(0.0)
    assert result["overall_verdict"] == "accept"


def test_hundred_percent_when_everything_flagged():
    chunks = [
        {"word_count": 100, "ai_probability": 0.9},
        {"word_count": 100, "ai_probability": 0.95},
    ]
    result = aggregate_chunk_results(chunks, chunk_probability_threshold=0.5, max_ai_percentage=15.0)
    assert result["ai_generated_percentage"] == pytest.approx(100.0)
    assert result["overall_verdict"] == "reject"


def test_rejects_empty_chunks():
    with pytest.raises(ValueError, match="must not be empty"):
        aggregate_chunk_results([])


def test_rejects_zero_total_words():
    with pytest.raises(ValueError, match="greater than zero"):
        aggregate_chunk_results([{"word_count": 0, "ai_probability": 0.9}])


# ── run_pipeline — orchestration, real chunking + mocked scorer ────

def _make_docx(path: str, text: str) -> str:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def test_pipeline_end_to_end_all_human(tmp_path):
    text = " ".join(["word"] * 50)
    path = _make_docx(str(tmp_path / "paper.docx"), text)

    fake_scorer = lambda t: {"ai_probability": 0.1}  # noqa: E731
    result = run_pipeline(path, words_per_chunk=20, scorer=fake_scorer)

    assert result["status"] == "complete"
    assert result["ai_generated_percentage"] == pytest.approx(0.0)
    assert result["overall_verdict"] == "accept"
    assert result["flagged_chunks"] == []


def test_pipeline_end_to_end_all_ai(tmp_path):
    text = " ".join(["word"] * 50)
    path = _make_docx(str(tmp_path / "paper.docx"), text)

    fake_scorer = lambda t: {"ai_probability": 0.9}  # noqa: E731
    result = run_pipeline(path, words_per_chunk=20, scorer=fake_scorer)

    assert result["ai_generated_percentage"] == pytest.approx(100.0)
    assert result["overall_verdict"] == "reject"
    assert result["flagged_chunk_count"] == 3


def test_pipeline_respects_organizer_configured_max_percentage(tmp_path):
    # 100 words, one 20-word chunk flagged (20%) — rejected at a strict
    # 15% policy, but would ACCEPT under a looser 25% policy. Confirms the
    # organizer's threshold is actually respected end-to-end, not hardcoded.
    text = "AIWORD " * 20 + "human " * 80

    def fake_scorer(t):
        return {"ai_probability": 0.9 if "AIWORD" in t else 0.1}

    path = _make_docx(str(tmp_path / "paper.docx"), text)

    strict = run_pipeline(path, words_per_chunk=20, scorer=fake_scorer, max_ai_percentage=15.0)
    assert strict["overall_verdict"] == "reject"

    lenient = run_pipeline(path, words_per_chunk=20, scorer=fake_scorer, max_ai_percentage=25.0)
    assert lenient["overall_verdict"] == "accept"


def test_pipeline_flagged_chunks_include_word_count_for_highlighting(tmp_path):
    text = "human words here " * 5 + "AIWORD AIWORD AIWORD " * 5 + "more human words here " * 5

    def fake_scorer(t):
        return {"ai_probability": 0.9 if "AIWORD" in t else 0.1}

    path = _make_docx(str(tmp_path / "paper.docx"), text)
    result = run_pipeline(path, words_per_chunk=15, scorer=fake_scorer)

    assert result["flagged_chunk_count"] == 1
    flagged = result["flagged_chunks"][0]
    assert "AIWORD" in flagged["text"]
    assert flagged["word_count"] == 15
    assert "start_char" in flagged


def test_pipeline_returns_error_for_missing_file(tmp_path):
    result = run_pipeline(str(tmp_path / "does_not_exist.docx"))
    assert result["status"] == "error"
    assert "Could not extract text" in result["error"]


def test_pipeline_returns_error_for_empty_document(tmp_path):
    doc = Document()
    path = str(tmp_path / "empty.docx")
    doc.save(path)
    result = run_pipeline(path)
    assert result["status"] == "error"
    assert "No extractable text" in result["error"]


def test_pipeline_returns_error_when_scorer_raises(tmp_path):
    text = " ".join(["word"] * 30)
    path = _make_docx(str(tmp_path / "paper.docx"), text)

    def failing_scorer(t):
        raise RuntimeError("CUDA out of memory")

    result = run_pipeline(path, scorer=failing_scorer)
    assert result["status"] == "error"
    assert "Scoring failed" in result["error"]


def test_pipeline_uses_followsci_by_default_when_no_scorer_given(tmp_path):
    text = " ".join(["word"] * 30)
    path = _make_docx(str(tmp_path / "paper.docx"), text)

    with patch("app.ai.followsci_check._score_text", return_value={"ai_probability": 0.9}) as mock_scorer:
        result = run_pipeline(path)

    assert mock_scorer.called
    assert result["ai_generated_percentage"] == pytest.approx(100.0)


# ── run_pipeline — highlighting wiring (update41) ──

def test_docx_flagged_chunks_have_empty_highlight_boxes_with_no_pdf(tmp_path):
    """The pre-update41 default: a .docx with no pdf_path_for_highlighting
    given at all — highlighting is skipped entirely, but the key is still
    present (empty list), not missing, so callers never have to branch on
    a maybe-absent key."""
    text = "human words here " * 5 + "AIWORD AIWORD AIWORD " * 5 + "more human words here " * 5

    def fake_scorer(t):
        return {"ai_probability": 0.9 if "AIWORD" in t else 0.1}

    path = _make_docx(str(tmp_path / "paper.docx"), text)
    result = run_pipeline(path, words_per_chunk=15, scorer=fake_scorer)

    assert result["flagged_chunk_count"] == 1
    assert result["flagged_chunks"][0]["highlight_boxes"] == []


def test_docx_with_failed_conversion_degrades_to_empty_highlight_boxes(tmp_path):
    """pdf_path_for_highlighting given but pointing at a non-.pdf path (the
    real-world shape of a failed Word->PDF conversion, where
    submissions.py's converted_pdf_path ends up None and falls back to the
    original file_path) — must not crash, must not attempt PyMuPDF at all
    since the extension check catches this before that point."""
    text = "human words here " * 5 + "AIWORD AIWORD AIWORD " * 5 + "more human words here " * 5

    def fake_scorer(t):
        return {"ai_probability": 0.9 if "AIWORD" in t else 0.1}

    path = _make_docx(str(tmp_path / "paper.docx"), text)
    # Simulates converted_pdf_path being None -> submissions.py's fallback
    # `converted_pdf_path or file_path` resolves to the .docx itself.
    result = run_pipeline(path, words_per_chunk=15, scorer=fake_scorer, pdf_path_for_highlighting=path)

    assert result["flagged_chunk_count"] == 1
    assert result["flagged_chunks"][0]["highlight_boxes"] == []


def test_real_pdf_extraction_used_when_pdf_path_for_highlighting_is_pdf(tmp_path):
    """When a real converted PDF is given, extraction must run against
    THAT file (for a real page_map), not the original .docx — confirmed by
    mocking extract_text and checking which path it was actually called
    with. This is the core of the update41 fix: previously page_map was
    always None for .docx regardless of what PDF was available."""
    text = " ".join(["word"] * 30)
    docx_path = _make_docx(str(tmp_path / "paper.docx"), text)
    fake_pdf_path = str(tmp_path / "paper.pdf")
    with open(fake_pdf_path, "wb") as f:
        f.write(b"%PDF-1.4 fake, never actually opened since extract_text is mocked below")

    fake_scorer = lambda t: {"ai_probability": 0.1}  # noqa: E731

    with patch("app.ai.ai_content_pipeline.extract_text") as mock_extract:
        mock_extract.return_value = (text, [(0, len(text), 1)])
        result = run_pipeline(docx_path, scorer=fake_scorer, pdf_path_for_highlighting=fake_pdf_path)

    mock_extract.assert_called_once_with(fake_pdf_path)
    assert result["status"] == "complete"


def test_compute_highlight_boxes_called_when_flagged_chunks_and_real_pdf_present(tmp_path):
    """compute_highlight_boxes should actually be invoked (and its result
    used) when there's both a flagged chunk and a usable PDF + page_map —
    mocked here since real PyMuPDF search_for() isn't testable in this
    environment (see ai_text_highlighting.py's own docstring)."""
    text = "human words here " * 5 + "AIWORD AIWORD AIWORD " * 5 + "more human words here " * 5
    docx_path = _make_docx(str(tmp_path / "paper.docx"), text)
    fake_pdf_path = str(tmp_path / "paper.pdf")
    with open(fake_pdf_path, "wb") as f:
        f.write(b"%PDF-1.4 fake")

    def fake_scorer(t):
        return {"ai_probability": 0.9 if "AIWORD" in t else 0.1}

    fake_boxes_result = "SENTINEL_RETURN_VALUE"

    with patch("app.ai.ai_content_pipeline.extract_text") as mock_extract, \
         patch("app.ai.ai_text_highlighting.compute_highlight_boxes") as mock_compute:
        mock_extract.return_value = (text, [(0, len(text), 1)])
        mock_compute.return_value = fake_boxes_result
        result = run_pipeline(docx_path, words_per_chunk=15, scorer=fake_scorer, pdf_path_for_highlighting=fake_pdf_path)

    assert mock_compute.called
    assert result["flagged_chunks"] == fake_boxes_result


def test_highlighting_failure_does_not_break_the_whole_check(tmp_path):
    """compute_highlight_boxes raising (corrupt PDF, PyMuPDF error, etc.)
    must degrade to empty highlight_boxes, not fail the entire AI-text
    check — the detection result is the primary output; highlighting is a
    bonus on top of it."""
    text = "human words here " * 5 + "AIWORD AIWORD AIWORD " * 5 + "more human words here " * 5
    docx_path = _make_docx(str(tmp_path / "paper.docx"), text)
    fake_pdf_path = str(tmp_path / "paper.pdf")
    with open(fake_pdf_path, "wb") as f:
        f.write(b"%PDF-1.4 fake")

    def fake_scorer(t):
        return {"ai_probability": 0.9 if "AIWORD" in t else 0.1}

    with patch("app.ai.ai_content_pipeline.extract_text") as mock_extract, \
         patch("app.ai.ai_text_highlighting.compute_highlight_boxes") as mock_compute:
        mock_extract.return_value = (text, [(0, len(text), 1)])
        mock_compute.side_effect = RuntimeError("simulated PyMuPDF failure")
        result = run_pipeline(docx_path, words_per_chunk=15, scorer=fake_scorer, pdf_path_for_highlighting=fake_pdf_path)

    assert result["status"] == "complete"
    assert result["flagged_chunk_count"] == 1
    assert result["flagged_chunks"][0]["highlight_boxes"] == []
