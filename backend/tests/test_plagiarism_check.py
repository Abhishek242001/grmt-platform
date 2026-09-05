from docx import Document

from app.ai.plagiarism_check import run_plagiarism_check

_IOT_PAPER = (
    "Industrial IoT deployments generate high-volume sensor streams that "
    "must be monitored for anomalies in real time. This paper presents a "
    "lightweight detection approach combining a sliding-window statistical "
    "baseline with a small gradient-boosted classifier, evaluated on three "
    "months of production line data from a mid-size manufacturing facility. "
    "Existing anomaly detection approaches often assume cloud connectivity, "
    "which is not always available on factory floors, motivating an "
    "edge-deployable alternative that trades some accuracy for much lower "
    "latency and no dependency on a persistent network connection."
)

_UNRELATED_PAPER = (
    "Drought stress significantly reduces crop yield in arid and semi-arid "
    "agricultural regions worldwide. This study examines the physiological "
    "response of three wheat cultivars to induced water deficit conditions "
    "during the critical grain-filling growth stage. Stomatal conductance, "
    "leaf water potential, and chlorophyll fluorescence were measured at "
    "regular intervals across a twelve-week greenhouse trial. Results "
    "indicate substantial variation in drought tolerance mechanisms between "
    "cultivars, with implications for future breeding programs targeting "
    "climate resilience in staple grain crops."
)

# A real ABSTRACT heading, matching a genuine paper's structure — needed
# for the external-check tests specifically (update46), since
# extract_abstract_and_conclusion() requires a real ABSTRACT section to
# find. _IOT_PAPER (above) deliberately has no such heading, for the tests
# where that's exactly the point (empty-candidates, missing-file, etc. —
# tests unrelated to external-check scoping).
_IOT_PAPER_WITH_ABSTRACT = (
    "ABSTRACT " + _IOT_PAPER + "\n"
    "I. INTRODUCTION\n"
    "This body text should never be sent to the external provider — only "
    "the abstract above should be, per update46's credit-conservation scoping."
)


def _make_docx(path: str, text: str) -> str:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(path)
    return path


def test_flags_near_identical_prior_submission(tmp_path):
    path = _make_docx(str(tmp_path / "paper.docx"), _IOT_PAPER)
    result = run_plagiarism_check(path, candidates=[{"submission_id": "prior-1", "text": _IOT_PAPER}])

    assert result["status"] == "complete"
    assert result["highest_similarity"] > 0.99
    assert len(result["matches"]) == 1
    assert len(result["issues"]) == 1
    assert "prior-1" in result["issues"][0]
    assert result["score"] < 5.0  # score is (1 - similarity)*100, near-identical -> near-zero score


def test_no_flag_for_unrelated_prior_submissions(tmp_path):
    path = _make_docx(str(tmp_path / "paper.docx"), _IOT_PAPER)
    result = run_plagiarism_check(path, candidates=[{"submission_id": "unrelated-1", "text": _UNRELATED_PAPER}])

    assert result["status"] == "complete"
    assert result["matches"] == []
    assert result["issues"] == []
    assert result["score"] > 90.0  # low similarity -> high (good) score


def test_no_candidates_is_a_clean_pass_not_an_error(tmp_path):
    """A submission with no prior submissions to compare against at all —
    e.g. the very first real submission on the whole platform — must not
    error just because there's nothing yet to compare it to."""
    path = _make_docx(str(tmp_path / "paper.docx"), _IOT_PAPER)
    result = run_plagiarism_check(path, candidates=[])

    assert result["status"] == "complete"
    assert result["score"] == 100.0
    assert result["matches"] == []
    assert result["candidates_compared"] == 0


def test_returns_error_for_missing_file(tmp_path):
    result = run_plagiarism_check(str(tmp_path / "does_not_exist.docx"), candidates=[])
    assert result["status"] == "error"
    assert "Could not extract text" in result["error"]


def test_returns_error_for_empty_document(tmp_path):
    doc = Document()
    path = str(tmp_path / "empty.docx")
    doc.save(path)
    result = run_plagiarism_check(path, candidates=[])
    assert result["status"] == "error"
    assert "No extractable text" in result["error"]


def test_returns_error_for_too_short_submitted_text(tmp_path):
    path = _make_docx(str(tmp_path / "short.docx"), "Way too short to compare.")
    result = run_plagiarism_check(path, candidates=[{"submission_id": "prior-1", "text": _IOT_PAPER}])
    assert result["status"] == "error"
    assert "too short" in result["error"].lower()


def test_custom_threshold_is_passed_through(tmp_path):
    path = _make_docx(str(tmp_path / "paper.docx"), _IOT_PAPER)
    # A near-identical candidate WOULD normally flag at the default
    # threshold — confirm a stricter custom threshold correctly suppresses it.
    result = run_plagiarism_check(
        path, candidates=[{"submission_id": "prior-1", "text": _IOT_PAPER}], flag_threshold=1.5
    )
    assert result["matches"] == []


# ── external_check_fn integration (update45) ──

def test_no_external_check_fn_means_external_key_is_none(tmp_path):
    """Default behavior (no external provider configured) — "external"
    must be present as None, not a missing key, so callers never have to
    branch on a maybe-absent key."""
    path = _make_docx(str(tmp_path / "paper.docx"), _IOT_PAPER)
    result = run_plagiarism_check(path, candidates=[])
    assert result["external"] is None


def test_external_check_fn_result_included_when_given(tmp_path):
    path = _make_docx(str(tmp_path / "paper.docx"), _IOT_PAPER_WITH_ABSTRACT)

    def fake_external_check(text):
        return {"status": "complete", "overall_similarity_pct": 0.0, "matches": []}

    result = run_plagiarism_check(path, candidates=[], external_check_fn=fake_external_check)
    assert result["external"]["status"] == "complete"


def test_external_check_fn_receives_only_the_abstract_not_full_document():
    """update46: confirms the credit-conservation scoping actually works —
    the callable receives ONLY the extracted abstract text, not the whole
    document (the introduction sentence in the fixture must NOT appear in
    what the external check received)."""
    received = {}

    def fake_external_check(text):
        received["text"] = text
        return {"status": "complete", "overall_similarity_pct": 0.0, "matches": []}

    import tempfile

    with tempfile.TemporaryDirectory() as d:
        path = _make_docx(f"{d}/paper.docx", _IOT_PAPER_WITH_ABSTRACT)
        run_plagiarism_check(path, candidates=[], external_check_fn=fake_external_check)

    assert _IOT_PAPER.strip() in received["text"]  # the real abstract content is there
    assert "should never be sent" not in received["text"]  # the introduction body text is NOT


def test_external_check_skipped_with_clear_reason_when_no_abstract_found():
    """A document with no real ABSTRACT heading (e.g. _IOT_PAPER, which
    deliberately has none) must not silently fall back to sending the
    whole document — that would defeat the entire point of the
    credit-conservation limit. external_check_fn must not even be called."""
    called = {"was_called": False}

    def fake_external_check(text):
        called["was_called"] = True
        return {"status": "complete", "overall_similarity_pct": 0.0, "matches": []}

    import tempfile

    with tempfile.TemporaryDirectory() as d:
        path = _make_docx(f"{d}/paper.docx", _IOT_PAPER)  # no ABSTRACT heading
        result = run_plagiarism_check(path, candidates=[], external_check_fn=fake_external_check)

    assert called["was_called"] is False
    assert result["external"]["status"] == "error"
    assert "No ABSTRACT section found" in result["external"]["error"]


def test_external_scan_full_document_flag_sends_whole_text_when_enabled():
    """Confirms the EXTERNAL_SCAN_FULL_DOCUMENT escape hatch actually works
    — flipping it sends the full document again, matching update45's
    original (pre-update46) behavior, for whenever there's budget to lift
    the abstract-only scoping."""
    import app.ai.plagiarism_check as plagiarism_check_module

    received = {}

    def fake_external_check(text):
        received["text"] = text
        return {"status": "complete", "overall_similarity_pct": 0.0, "matches": []}

    import tempfile

    original_flag = plagiarism_check_module.EXTERNAL_SCAN_FULL_DOCUMENT
    try:
        plagiarism_check_module.EXTERNAL_SCAN_FULL_DOCUMENT = True
        with tempfile.TemporaryDirectory() as d:
            path = _make_docx(f"{d}/paper.docx", _IOT_PAPER_WITH_ABSTRACT)
            run_plagiarism_check(path, candidates=[], external_check_fn=fake_external_check)
        assert "should never be sent" in received["text"]  # now the intro body text IS included
    finally:
        plagiarism_check_module.EXTERNAL_SCAN_FULL_DOCUMENT = original_flag


def test_external_matches_added_to_issues():
    def fake_external_check(text):
        return {
            "status": "complete",
            "overall_similarity_pct": 40.0,
            "matches": [{"source_title": "A Real External Paper", "source_url": "https://example.com/paper", "similarity_pct": 40.0}],
        }

    import tempfile

    with tempfile.TemporaryDirectory() as d:
        path = _make_docx(f"{d}/paper.docx", _IOT_PAPER_WITH_ABSTRACT)
        result = run_plagiarism_check(path, candidates=[], external_check_fn=fake_external_check)

    assert len(result["issues"]) == 1
    assert "A Real External Paper" in result["issues"][0]
    assert "40.0%" in result["issues"][0]


def test_external_high_similarity_lowers_combined_score():
    """A clean self-submission comparison (no candidates at all -> score
    100) must NOT mask a real external match — the combined score should
    reflect the worse of the two."""
    def fake_external_check(text):
        return {"status": "complete", "overall_similarity_pct": 60.0, "matches": []}

    import tempfile

    with tempfile.TemporaryDirectory() as d:
        path = _make_docx(f"{d}/paper.docx", _IOT_PAPER_WITH_ABSTRACT)
        result = run_plagiarism_check(path, candidates=[], external_check_fn=fake_external_check)

    # self-submission alone would score 100 (no candidates); external
    # found 60% similarity -> external_score = 40 -> combined must be 40.
    assert result["score"] == 40.0


def test_external_check_fn_raising_does_not_break_self_submission_result():
    """An unhandled exception from the external check (network failure,
    bug, whatever) must not take down the whole plagiarism check — the
    self-submission comparison already succeeded and is still worth
    reporting."""
    def broken_external_check(text):
        raise RuntimeError("simulated external provider crash")

    import tempfile

    with tempfile.TemporaryDirectory() as d:
        path = _make_docx(f"{d}/paper.docx", _IOT_PAPER_WITH_ABSTRACT)
        result = run_plagiarism_check(path, candidates=[], external_check_fn=broken_external_check)

    assert result["status"] == "complete"  # the overall check still succeeded
    assert result["external"]["status"] == "error"
    assert "simulated external provider crash" in result["external"]["error"]


def test_external_check_own_error_status_does_not_add_fake_issues():
    """When the external provider itself reports an error (e.g. out of
    credits), no bogus "X% similar" issues should be synthesized from it."""
    def failing_external_check(text):
        return {"status": "error", "error": "Insufficient credits", "error_code": "PAYMENT_REQUIRED"}

    import tempfile

    with tempfile.TemporaryDirectory() as d:
        path = _make_docx(f"{d}/paper.docx", _IOT_PAPER_WITH_ABSTRACT)
        result = run_plagiarism_check(path, candidates=[], external_check_fn=failing_external_check)

    assert result["issues"] == []
    assert result["external"]["status"] == "error"
