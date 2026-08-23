from unittest.mock import patch

from app.ai.followsci_check import run_ai_text_detection_check, verdict


def test_verdict_above_threshold_is_likely_ai():
    assert verdict(0.8, threshold=0.5) == "likely_ai"


def test_verdict_below_threshold_is_likely_human():
    assert verdict(0.2, threshold=0.5) == "likely_human"


def test_verdict_exactly_at_threshold_is_likely_human():
    assert verdict(0.5, threshold=0.5) == "likely_human"


def test_returns_error_for_missing_file(tmp_path):
    result = run_ai_text_detection_check(str(tmp_path / "does_not_exist.docx"))
    assert result["status"] == "error"
    assert "Could not extract text" in result["error"]


def test_returns_error_for_empty_document(tmp_path):
    from docx import Document

    doc = Document()
    path = str(tmp_path / "empty.docx")
    doc.save(path)
    result = run_ai_text_detection_check(path)
    assert result["status"] == "error"
    assert "No extractable text" in result["error"]


def test_returns_error_when_scoring_fails(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Some real text content here.")
    path = str(tmp_path / "paper.docx")
    doc.save(path)

    with patch("app.ai.followsci_check._score_text", side_effect=RuntimeError("CUDA out of memory")):
        result = run_ai_text_detection_check(path)

    assert result["status"] == "error"
    assert "Scoring failed" in result["error"]


def test_happy_path_shape_with_mocked_ai_result(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Some real text content here.")
    path = str(tmp_path / "paper.docx")
    doc.save(path)

    with patch("app.ai.followsci_check._score_text", return_value={"ai_probability": 0.95}):
        result = run_ai_text_detection_check(path)

    assert result["status"] == "complete"
    assert result["verdict"] == "likely_ai"
    assert len(result["issues"]) == 1


def test_happy_path_shape_with_mocked_human_result(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Some real text content here.")
    path = str(tmp_path / "paper.docx")
    doc.save(path)

    with patch("app.ai.followsci_check._score_text", return_value={"ai_probability": 0.05}):
        result = run_ai_text_detection_check(path)

    assert result["status"] == "complete"
    assert result["verdict"] == "likely_human"
    assert result["issues"] == []
