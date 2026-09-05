import json
from unittest.mock import patch

import httpx
from docx import Document

from app.ai.logical_consistency_check import run_logical_consistency_check


def _make_docx_with_sections(path: str, abstract: str, conclusion: str) -> str:
    doc = Document()
    doc.add_paragraph("PAPER TITLE")
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph(abstract)
    doc.add_paragraph("INTRODUCTION")
    doc.add_paragraph("Some intro text that should be ignored by this check.")
    doc.add_paragraph("CONCLUSION")
    doc.add_paragraph(conclusion)
    doc.add_paragraph("REFERENCES")
    doc.add_paragraph("[1] Some reference.")
    doc.save(path)
    return path


# ── section-extraction failure — real logic, no mocking needed ─────

def test_returns_error_when_no_abstract_or_conclusion_sections(tmp_path):
    doc = Document()
    doc.add_paragraph("Just some plain text with no section headings.")
    path = str(tmp_path / "paper.docx")
    doc.save(path)

    result = run_logical_consistency_check(path)
    assert result["status"] == "error"
    assert "abstract and conclusion" in result["error"]


def test_returns_error_for_missing_file(tmp_path):
    result = run_logical_consistency_check(str(tmp_path / "does_not_exist.docx"))
    assert result["status"] == "error"
    assert "Could not extract text" in result["error"]


# ── Ollama call — mocked, since a real Ollama service isn't available here ──

def test_returns_error_when_ollama_unreachable(tmp_path):
    path = _make_docx_with_sections(str(tmp_path / "paper.docx"), "We achieve 95% accuracy.", "We achieved good results.")

    with patch("app.ai.logical_consistency_check._call_ollama", side_effect=httpx.RequestError("Connection refused")):
        result = run_logical_consistency_check(path)

    assert result["status"] == "error"
    assert "Could not reach Ollama" in result["error"]


def test_returns_error_on_ollama_http_error(tmp_path):
    path = _make_docx_with_sections(str(tmp_path / "paper.docx"), "We achieve 95% accuracy.", "We achieved good results.")

    mock_response = httpx.Response(500, request=httpx.Request("POST", "http://localhost:11434/api/chat"))
    with patch(
        "app.ai.logical_consistency_check._call_ollama",
        side_effect=httpx.HTTPStatusError("Server error", request=mock_response.request, response=mock_response),
    ):
        result = run_logical_consistency_check(path)

    assert result["status"] == "error"
    assert "Ollama returned an error" in result["error"]
    assert "500" in result["error"]


def test_full_pipeline_consistent_case(tmp_path):
    path = _make_docx_with_sections(
        str(tmp_path / "paper.docx"),
        "Our method achieves 95% accuracy on the benchmark.",
        "We demonstrated 95% accuracy, confirming our approach is effective.",
    )
    mocked_response = json.dumps({"consistent": True, "findings": []})

    with patch("app.ai.logical_consistency_check._call_ollama", return_value=mocked_response):
        result = run_logical_consistency_check(path)

    assert result["status"] == "complete"
    assert result["consistent"] is True
    assert result["score"] == 100.0
    assert result["issues"] == []


def test_full_pipeline_inconsistent_case(tmp_path):
    path = _make_docx_with_sections(
        str(tmp_path / "paper.docx"),
        "Our method achieves 95% accuracy on the benchmark.",
        "We achieved approximately 80% accuracy in our final evaluation.",
    )
    mocked_response = json.dumps({
        "consistent": False,
        "findings": [{
            "abstract_claim": "achieves 95% accuracy",
            "conclusion_statement": "achieved approximately 80% accuracy",
            "explanation": "The accuracy figure differs by 15 percentage points between the abstract and conclusion.",
        }],
    })

    with patch("app.ai.logical_consistency_check._call_ollama", return_value=mocked_response):
        result = run_logical_consistency_check(path)

    assert result["status"] == "complete"
    assert result["consistent"] is False
    assert result["score"] == 0.0
    assert len(result["issues"]) == 1
    assert "95% accuracy" in result["issues"][0]
    assert "80% accuracy" in result["issues"][0]


def test_returns_error_when_ollama_response_fails_validation(tmp_path):
    path = _make_docx_with_sections(str(tmp_path / "paper.docx"), "Claim.", "Statement.")

    with patch("app.ai.logical_consistency_check._call_ollama", return_value="not valid json at all"):
        result = run_logical_consistency_check(path)

    assert result["status"] == "error"
    assert "Could not parse Ollama response" in result["error"]


def test_call_ollama_receives_the_extracted_sections_not_the_whole_document(tmp_path):
    """Confirms the check sends ONLY the abstract/conclusion to the model,
    not the full document text (including the deliberately-irrelevant
    introduction paragraph) — keeps the prompt focused and avoids wasting
    context on sections this check doesn't reason about."""
    path = _make_docx_with_sections(
        str(tmp_path / "paper.docx"),
        "UNIQUE_ABSTRACT_MARKER claim here.",
        "UNIQUE_CONCLUSION_MARKER statement here.",
    )

    captured = {}

    def fake_call(abstract, conclusion):
        captured["abstract"] = abstract
        captured["conclusion"] = conclusion
        return json.dumps({"consistent": True, "findings": []})

    with patch("app.ai.logical_consistency_check._call_ollama", side_effect=fake_call):
        run_logical_consistency_check(path)

    assert "UNIQUE_ABSTRACT_MARKER" in captured["abstract"]
    assert "UNIQUE_CONCLUSION_MARKER" in captured["conclusion"]
    assert "UNIQUE_ABSTRACT_MARKER" not in captured["conclusion"]
    assert "intro text" not in captured["abstract"]
    assert "intro text" not in captured["conclusion"]
