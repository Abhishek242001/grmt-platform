from unittest.mock import patch

import pytest

from app.ai.radar_check import run_ai_text_detection_check, verdict


# ── verdict — boundary behavior ─────────────────────────────────

def test_verdict_above_threshold_is_likely_ai():
    assert verdict(0.8, threshold=0.5) == "likely_ai"


def test_verdict_below_threshold_is_likely_human():
    assert verdict(0.2, threshold=0.5) == "likely_human"


def test_verdict_exactly_at_threshold_is_likely_human():
    assert verdict(0.5, threshold=0.5) == "likely_human"


# ── run_ai_text_detection_check — error handling, mocking the model call ──
#
# _score_text() itself needs a real GPU + the real RADAR model, which isn't
# available in this environment — same constraint as the other two checks.
# What CAN be tested here without either: the orchestration logic around
# it (extraction failures, empty documents, scoring failures, and the
# happy-path response shape), by mocking _score_text the same way
# grammar_check's tests mock the LanguageTool HTTP call.

def test_returns_error_for_missing_file(tmp_path):
    result = run_ai_text_detection_check(str(tmp_path / "does_not_exist.docx"))
    assert result["status"] == "error"
    assert "Could not extract text" in result["error"]


def test_returns_error_for_empty_document(tmp_path):
    from docx import Document

    doc = Document()
    path = str(tmp_path / "empty.docx")
    doc.save(path)
    result = run_ai_text_detection_check(path)
    assert result["status"] == "error"
    assert "No extractable text" in result["error"]


def test_returns_error_when_scoring_fails(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Some real text content here.")
    path = str(tmp_path / "paper.docx")
    doc.save(path)

    with patch("app.ai.radar_check._score_text", side_effect=RuntimeError("CUDA out of memory")):
        result = run_ai_text_detection_check(path)

    assert result["status"] == "error"
    assert "Scoring failed" in result["error"]
    assert "CUDA out of memory" in result["error"]


def test_happy_path_shape_with_mocked_ai_result(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Some real text content here.")
    path = str(tmp_path / "paper.docx")
    doc.save(path)

    with patch("app.ai.radar_check._score_text", return_value={"ai_probability": 0.92}):
        result = run_ai_text_detection_check(path)

    assert result["status"] == "complete"
    assert result["ai_probability"] == 0.92
    assert result["verdict"] == "likely_ai"
    assert len(result["issues"]) == 1
    assert "likely AI-generated" in result["issues"][0]


def test_happy_path_shape_with_mocked_human_result(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Some real text content here.")
    path = str(tmp_path / "paper.docx")
    doc.save(path)

    with patch("app.ai.radar_check._score_text", return_value={"ai_probability": 0.08}):
        result = run_ai_text_detection_check(path)

    assert result["status"] == "complete"
    assert result["verdict"] == "likely_human"
    assert result["issues"] == []


def test_custom_threshold_is_respected(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Some real text content here.")
    path = str(tmp_path / "paper.docx")
    doc.save(path)

    # 0.6 probability would be "likely_ai" at the default 0.5 threshold,
    # but "likely_human" at a stricter custom 0.7 threshold.
    with patch("app.ai.radar_check._score_text", return_value={"ai_probability": 0.6}):
        result = run_ai_text_detection_check(path, threshold=0.7)

    assert result["verdict"] == "likely_human"
