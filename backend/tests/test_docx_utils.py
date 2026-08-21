import io
import zipfile

import pytest
from docx import Document

from app.ai.docx_utils import extract_textbox_paragraphs, open_docx


def _make_legacy_namespace_docx(path: str) -> str:
    """Builds a real, valid .docx via python-docx, then rewrites its XML to
    use the alternate 'purl.oclc.org' namespace instead of the canonical
    'schemas.openxmlformats.org' one — reproducing the exact real-world
    pattern found in a genuine third-party-generated template (planning log
    §33), so this test exercises the real failure mode, not an invented one."""
    doc = Document()
    doc.add_paragraph("Body text in a legacy-namespace document.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    output = io.BytesIO()
    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith((".xml", ".rels")):
                data = data.replace(
                    b"http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                    b"http://purl.oclc.org/ooxml/wordprocessingml/main",
                )
                data = data.replace(
                    b"http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                    b"http://purl.oclc.org/ooxml/officeDocument/relationships",
                )
            zout.writestr(item, data)

    with open(path, "wb") as f:
        f.write(output.getvalue())
    return path


def test_genuine_word_docx_opens_without_needing_the_fallback():
    """The common case — a real python-docx-generated file — must open on
    the first, fast path without ever needing normalization."""
    doc = Document()
    doc.add_paragraph("Ordinary body text.")
    path = "/tmp/_test_genuine_docx.docx"
    doc.save(path)

    opened = open_docx(path)
    assert opened.paragraphs[0].text == "Ordinary body text."


def test_standard_document_open_rejects_legacy_namespace_file():
    """Confirms the failure mode this fix addresses actually exists — a
    document.Document() call on a legacy-namespace file must fail the way
    the real file did, or this whole test file is testing the wrong thing."""
    path = _make_legacy_namespace_docx("/tmp/_test_legacy_ns_raw.docx")
    with pytest.raises(KeyError):
        Document(path)


def test_open_docx_falls_back_and_reads_legacy_namespace_file_correctly():
    path = _make_legacy_namespace_docx("/tmp/_test_legacy_ns.docx")
    doc = open_docx(path)
    assert doc.paragraphs[0].text == "Body text in a legacy-namespace document."


# ── extract_textbox_paragraphs ──────────────────────────────────
#
# Reproduces a genuine real-world failure: a real IEEE-format submission's
# figure and table captions were inserted via a text box (exactly as IEEE's
# own template guidance recommends — "use a text box... more stable than
# directly inserting a picture"), and were completely invisible to every
# check reading document text, because doc.paragraphs never descends into
# a text box's nested w:txbxContent element. Confirmed empirically against
# the actual failure before writing this fix (table_figure_check flagged
# both a figure and a table as "referenced but no caption found" when the
# caption genuinely existed in the file).

def _add_legacy_vml_textbox(doc, text: str) -> None:
    """Appends a paragraph whose run contains a legacy VML text box (the
    `<v:textbox>` wrapper, as opposed to the modern DrawingML one) —
    reproduces the exact real-world structure IEEE's own template guidance
    describes, not an invented XML shape."""
    from lxml import etree

    p = doc.add_paragraph()
    run = p.add_run()
    textbox_xml = f"""
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml">
      <v:shape>
        <v:textbox>
          <w:txbxContent xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:p><w:r><w:t>{text}</w:t></w:r></w:p>
          </w:txbxContent>
        </v:textbox>
      </v:shape>
    </w:pict>
    """
    run._r.append(etree.fromstring(textbox_xml))


def test_extract_textbox_paragraphs_finds_legacy_vml_textbox_content(tmp_path):
    doc = Document()
    doc.add_paragraph("Ordinary body paragraph, outside any text box.")
    _add_legacy_vml_textbox(doc, "Fig. 1. Caption living inside a text box.")
    path = str(tmp_path / "textbox.docx")
    doc.save(path)

    opened = open_docx(path)
    # Confirms the gap this fix addresses actually exists — doc.paragraphs
    # must NOT see the text-box content, or this whole fix is solving a
    # problem that doesn't exist.
    assert all("Fig. 1" not in p.text for p in opened.paragraphs)

    textbox_paragraphs = extract_textbox_paragraphs(opened)
    assert textbox_paragraphs == ["Fig. 1. Caption living inside a text box."]


def test_extract_textbox_paragraphs_returns_empty_list_when_no_textboxes(tmp_path):
    doc = Document()
    doc.add_paragraph("Just an ordinary document with no text boxes at all.")
    path = str(tmp_path / "no_textbox.docx")
    doc.save(path)

    assert extract_textbox_paragraphs(open_docx(path)) == []
