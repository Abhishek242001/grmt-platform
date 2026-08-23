import io
import json

from docx import Document


def _signup(client, email, role="researcher"):
    r = client.post("/api/auth/signup", json={
        "email": email, "password": "Password1", "full_name": "Test User", "role": role,
    })
    return r.json()["access_token"]


def _auth(token):
    return {"Authorization": f"Bearer {token}"}


def _make_conference(client, org_token):
    r = client.post("/api/conferences", json={"name": "ICSE 2026"}, headers=_auth(org_token))
    return r.json()["id"]


def _submit(client, res_token, conf_id, title="Paper"):
    return client.post(
        "/api/submissions",
        json={
            "conference_id": conf_id, "title": title,
            "original_filename": "paper.docx", "original_file_url": "placeholder://uploads/paper.docx",
        },
        headers=_auth(res_token),
    )


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _make_pdf_bytes(text: str) -> bytes:
    import pymupdf
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((50, 72), text)
    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


# ── Existing coverage (ownership, visibility, resubmit) ──────────────────

def test_researcher_can_create_and_view_own_submission(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")

    r = _submit(client, res_token, conf_id, title="My Paper")
    assert r.status_code == 201
    sub_id = r.json()["id"]

    r = client.get(f"/api/submissions/{sub_id}", headers=_auth(res_token))
    assert r.status_code == 200


def test_researcher_cannot_view_another_researchers_submission(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res1 = _signup(client, "res1@example.com", role="researcher")
    res2 = _signup(client, "res2@example.com", role="researcher")

    r = _submit(client, res1, conf_id, title="My Paper")
    sub_id = r.json()["id"]

    r = client.get(f"/api/submissions/{sub_id}", headers=_auth(res2))
    assert r.status_code == 404


def test_researcher_cannot_view_another_researchers_ai_report(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res1 = _signup(client, "res1@example.com", role="researcher")
    res2 = _signup(client, "res2@example.com", role="researcher")

    r = _submit(client, res1, conf_id, title="My Paper")
    sub_id = r.json()["id"]

    r = client.get(f"/api/submissions/{sub_id}/ai-report", headers=_auth(res2))
    assert r.status_code == 404


def test_conference_organizer_can_view_submissions_to_their_conference(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")

    r = _submit(client, res_token, conf_id, title="My Paper")
    sub_id = r.json()["id"]

    r = client.get(f"/api/submissions/{sub_id}", headers=_auth(org_token))
    assert r.status_code == 200


def test_unrelated_organizer_cannot_view_submission(client):
    org1 = _signup(client, "org1@example.com", role="organizer")
    org2 = _signup(client, "org2@example.com", role="organizer")
    conf_id = _make_conference(client, org1)
    res_token = _signup(client, "res@example.com", role="researcher")

    r = _submit(client, res_token, conf_id, title="My Paper")
    sub_id = r.json()["id"]

    r = client.get(f"/api/submissions/{sub_id}", headers=_auth(org2))
    assert r.status_code == 404


def test_assigned_reviewer_can_view_submission(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    rev_token = _signup(client, "rev@example.com", role="reviewer")
    conf_id = _make_conference(client, org_token)
    client.post(f"/api/conferences/{conf_id}/reviewers", json={"email": "rev@example.com"}, headers=_auth(org_token))

    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id, title="Paper")
    sub_id = r.json()["id"]

    r = client.get(f"/api/submissions/{sub_id}", headers=_auth(rev_token))
    assert r.status_code == 200


def test_unassigned_reviewer_cannot_view_submission(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    rev_token = _signup(client, "rev@example.com", role="reviewer")
    conf_id = _make_conference(client, org_token)

    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id, title="Paper")
    sub_id = r.json()["id"]

    r = client.get(f"/api/submissions/{sub_id}", headers=_auth(rev_token))
    assert r.status_code == 404


def test_researcher_can_list_own_submissions(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")
    _submit(client, res_token, conf_id, title="Paper 1")
    _submit(client, res_token, conf_id, title="Paper 2")

    r = client.get("/api/submissions/mine", headers=_auth(res_token))
    assert r.status_code == 200
    assert len(r.json()) == 2


def test_create_submission_creates_initial_version(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")

    r = _submit(client, res_token, conf_id)
    assert r.status_code == 201
    sub_id = r.json()["id"]

    r = client.get(f"/api/submissions/{sub_id}/history", headers=_auth(res_token))
    assert r.status_code == 200
    assert len(r.json()) == 1
    assert r.json()[0]["version_number"] == 1


def test_resubmit_requires_revise_resubmit_status(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id)
    sub_id = r.json()["id"]

    r = client.post(
        f"/api/submissions/{sub_id}/resubmit",
        json={"original_filename": "v2.docx", "original_file_url": "placeholder://uploads/v2.docx"},
        headers=_auth(res_token),
    )
    assert r.status_code == 400


def test_other_researcher_cannot_resubmit_someone_elses_paper(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res1 = _signup(client, "res1@example.com", role="researcher")
    res2 = _signup(client, "res2@example.com", role="researcher")
    r = _submit(client, res1, conf_id)
    sub_id = r.json()["id"]

    r = client.post(
        f"/api/submissions/{sub_id}/resubmit",
        json={"original_filename": "v2.docx", "original_file_url": "placeholder://uploads/v2.docx"},
        headers=_auth(res2),
    )
    assert r.status_code == 404


def test_reviewer_sees_assigned_submissions_only(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    rev_token = _signup(client, "rev@example.com", role="reviewer")
    conf1 = _make_conference(client, org_token)
    conf2_r = client.post("/api/conferences", json={"name": "Other Conf"}, headers=_auth(org_token))
    conf2 = conf2_r.json()["id"]
    client.post(f"/api/conferences/{conf1}/reviewers", json={"email": "rev@example.com"}, headers=_auth(org_token))

    res_token = _signup(client, "res@example.com", role="researcher")
    _submit(client, res_token, conf1, title="In Assigned Conf")
    _submit(client, res_token, conf2, title="Not Assigned")

    r = client.get("/api/submissions/assigned", headers=_auth(rev_token))
    assert r.status_code == 200
    titles = [s["title"] for s in r.json()]
    assert titles == ["In Assigned Conf"]


# ── Phase 2: real file upload + grammar check orchestration ──────────────

def test_upload_rejects_unsupported_file_type(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id)
    sub_id = r.json()["id"]

    r = client.post(
        f"/api/submissions/{sub_id}/upload",
        files={"file": ("notes.txt", b"just some text", "text/plain")},
        headers=_auth(res_token),
    )
    assert r.status_code == 400


def test_other_researcher_cannot_upload_to_someone_elses_submission(client):
    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res1 = _signup(client, "res1@example.com", role="researcher")
    res2 = _signup(client, "res2@example.com", role="researcher")
    r = _submit(client, res1, conf_id)
    sub_id = r.json()["id"]

    docx_bytes = _make_docx_bytes(["This is fine."])
    r = client.post(
        f"/api/submissions/{sub_id}/upload",
        files={"file": ("paper.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=_auth(res2),
    )
    assert r.status_code == 404


def test_upload_runs_grammar_check_and_stores_ai_report(client, monkeypatch):
    """The real proof: upload a genuine .docx, have the grammar-check module hit
    a MOCKED LanguageTool response (no live LanguageTool in the test suite, per
    development_rule.md's own testing philosophy for external AI services), and
    confirm a real AIReport row lands with a real computed score."""

    def fake_post(url, data=None, timeout=None):
        class FakeResponse:
            def raise_for_status(self):
                pass

            def json(self):
                return {
                    "matches": [
                        {
                            "message": "Possible grammar issue",
                            "shortMessage": "Grammar",
                            "offset": 0,
                            "length": 4,
                            "rule": {"id": "TEST_RULE", "category": {"name": "Grammar"}},
                        }
                    ]
                }

        return FakeResponse()

    import app.ai.grammar_check as grammar_check_module
    monkeypatch.setattr(grammar_check_module.httpx, "post", fake_post)

    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id)
    sub_id = r.json()["id"]

    docx_bytes = _make_docx_bytes(["This paper have some grammar issue in it."])
    r = client.post(
        f"/api/submissions/{sub_id}/upload",
        files={"file": ("paper.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=_auth(res_token),
    )
    assert r.status_code == 200

    # BackgroundTasks run synchronously within TestClient's request lifecycle,
    # so by the time this line runs, the grammar check AND the gate-evaluation
    # engine have both already completed. No gate rule is configured for this
    # conference, so status should have moved to "in_human_review" — never
    # "ai_review_passed" (only 1 of 7 checks exists), and never stuck on
    # "processing" forever (the original gap this whole feature closes).
    sub_check = client.get(f"/api/submissions/{sub_id}", headers=_auth(res_token))
    assert sub_check.json()["status"] == "in_human_review"

    r = client.get(f"/api/submissions/{sub_id}/ai-report", headers=_auth(res_token))
    assert r.status_code == 200
    reports = r.json()
    # Four reports now — grammar, format compliance, table/figure
    # consistency, AND ai_text (AI-generated-content detection) all run
    # per upload.
    assert len(reports) == 4
    report_by_type = {r["check_type"]: r for r in reports}
    assert set(report_by_type.keys()) == {"grammar", "format", "table_figure", "ai_text"}
    assert report_by_type["grammar"]["status"] == "complete"
    assert report_by_type["format"]["status"] == "complete"
    assert report_by_type["table_figure"]["status"] == "complete"
    # ai_text needs torch/transformers + ideally a GPU — gracefully
    # degrades to status "error" (not a crash) in an environment without
    # them, and "complete" where they're available. Both are valid,
    # environment-dependent outcomes; what matters is the report exists.
    assert report_by_type["ai_text"]["status"] in ("complete", "error")

    result = json.loads(report_by_type["grammar"]["result_json"])
    assert result["error_count"] == 1
    assert result["score"] is not None
    assert result["matches"][0]["rule_id"] == "TEST_RULE"


def test_upload_with_configured_hard_gate_actually_hard_fails_submission(client, monkeypatch):
    """Full integration: real upload, real gate rule configured through the real
    API (not a synthetic DB row), a badly-scored mocked check, and confirmation
    the submission actually lands on ai_review_hard_failed — proving the upload
    endpoint, grammar check, and gate engine are correctly wired together, not
    just individually correct in isolation."""

    def fake_post_many_errors(url, data=None, timeout=None):
        class FakeResponse:
            def raise_for_status(self):
                pass

            def json(self):
                # Enough matches relative to a short doc to push the score well
                # below any reasonable threshold.
                return {
                    "matches": [
                        {"message": "issue", "shortMessage": "x", "offset": 0, "length": 1,
                         "rule": {"id": f"RULE_{i}", "category": {"name": "Grammar"}}}
                        for i in range(10)
                    ]
                }

        return FakeResponse()

    import app.ai.grammar_check as grammar_check_module
    monkeypatch.setattr(grammar_check_module.httpx, "post", fake_post_many_errors)

    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)

    # Configure a real hard gate on grammar, through the real API.
    r = client.put(
        f"/api/conferences/{conf_id}/gate-rules",
        json=[{"check_type": "grammar", "is_hard_gate": True, "threshold": 90}],
        headers=_auth(org_token),
    )
    assert r.status_code == 200

    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id)
    sub_id = r.json()["id"]

    docx_bytes = _make_docx_bytes(["Short document with many issues."])
    r = client.post(
        f"/api/submissions/{sub_id}/upload",
        files={"file": ("paper.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=_auth(res_token),
    )
    assert r.status_code == 200

    sub_check = client.get(f"/api/submissions/{sub_id}", headers=_auth(res_token))
    assert sub_check.json()["status"] == "ai_review_hard_failed"


def test_pdf_upload_also_runs_grammar_check(client, monkeypatch):
    """Confirms the PDF path works too, not just .docx — a real PDF (created via
    PyMuPDF itself, not a fake byte string) through the real upload endpoint,
    real column-aware extraction, mocked LanguageTool call."""

    def fake_post(url, data=None, timeout=None):
        class FakeResponse:
            def raise_for_status(self):
                pass

            def json(self):
                return {"matches": []}

        return FakeResponse()

    import app.ai.grammar_check as grammar_check_module
    monkeypatch.setattr(grammar_check_module.httpx, "post", fake_post)

    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id)
    sub_id = r.json()["id"]

    pdf_bytes = _make_pdf_bytes("This is a test PDF submission for grammar checking.")
    r = client.post(
        f"/api/submissions/{sub_id}/upload",
        files={"file": ("paper.pdf", pdf_bytes, "application/pdf")},
        headers=_auth(res_token),
    )
    assert r.status_code == 200

    r = client.get(f"/api/submissions/{sub_id}/ai-report", headers=_auth(res_token))
    reports = r.json()
    assert len(reports) == 4
    report_by_type = {r["check_type"]: r for r in reports}
    result = json.loads(report_by_type["grammar"]["result_json"])
    assert result["status"] == "complete"  # extraction + LanguageTool call both succeeded


def test_docx_upload_populates_converted_pdf_url(client, monkeypatch):
    """Word->PDF pipeline integration test — real LibreOffice conversion
    (not mocked; soffice is a real, testable dependency here), triggered by
    a real .docx upload through the real endpoint. Confirms the resulting
    converted_pdf_url is a genuinely valid, openable PDF, not just a
    non-null string."""

    def fake_post(url, data=None, timeout=None):
        class FakeResponse:
            def raise_for_status(self):
                pass

            def json(self):
                return {"matches": []}

        return FakeResponse()

    import app.ai.grammar_check as grammar_check_module
    monkeypatch.setattr(grammar_check_module.httpx, "post", fake_post)

    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id)
    sub_id = r.json()["id"]

    docx_bytes = _make_docx_bytes(["Body text for the word-to-pdf conversion test."])
    r = client.post(
        f"/api/submissions/{sub_id}/upload",
        files={"file": ("paper.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=_auth(res_token),
    )
    assert r.status_code == 200

    r = client.get(f"/api/submissions/{sub_id}/history", headers=_auth(res_token))
    assert r.status_code == 200
    versions = r.json()
    assert len(versions) == 1
    converted_path = versions[0]["converted_pdf_url"]
    assert converted_path is not None
    assert converted_path.endswith(".pdf")

    import pymupdf
    pdf = pymupdf.open(converted_path)
    assert len(pdf) >= 1
    assert "Body text for the word-to-pdf conversion test." in pdf[0].get_text()


def test_pdf_upload_sets_converted_pdf_url_to_itself(client, monkeypatch):
    """An already-PDF upload needs no conversion — converted_pdf_url should
    just point at the original file, so the frontend never has to branch
    on original file type to find a PDF to display."""

    def fake_post(url, data=None, timeout=None):
        class FakeResponse:
            def raise_for_status(self):
                pass

            def json(self):
                return {"matches": []}

        return FakeResponse()

    import app.ai.grammar_check as grammar_check_module
    monkeypatch.setattr(grammar_check_module.httpx, "post", fake_post)

    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id)
    sub_id = r.json()["id"]

    pdf_bytes = _make_pdf_bytes("Already a PDF, no conversion needed.")
    r = client.post(
        f"/api/submissions/{sub_id}/upload",
        files={"file": ("paper.pdf", pdf_bytes, "application/pdf")},
        headers=_auth(res_token),
    )
    assert r.status_code == 200

    r = client.get(f"/api/submissions/{sub_id}/history", headers=_auth(res_token))
    versions = r.json()
    assert versions[0]["converted_pdf_url"] is not None
    assert versions[0]["converted_pdf_url"].endswith("paper.pdf")


def test_upload_also_runs_format_compliance_check(client, monkeypatch):
    """Full integration: real upload, real format-compliance measurement
    against a genuinely well-formed IEEE .docx (correct margins, correct
    font size, real Abstract/References headings) — proving the check
    actually measures the real uploaded file, not a canned response."""

    def fake_post(url, data=None, timeout=None):
        class FakeResponse:
            def raise_for_status(self):
                pass

            def json(self):
                return {"matches": []}

        return FakeResponse()

    import app.ai.grammar_check as grammar_check_module
    monkeypatch.setattr(grammar_check_module.httpx, "post", fake_post)

    org_token = _signup(client, "org@example.com", role="organizer")
    conf_id = _make_conference(client, org_token)
    res_token = _signup(client, "res@example.com", role="researcher")
    r = _submit(client, res_token, conf_id)
    sub_id = r.json()["id"]

    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt

    doc = DocxDocument()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.125)
    section.left_margin = Inches(0.8125)
    section.right_margin = Inches(0.8125)
    p = doc.add_paragraph()
    run = p.add_run("Body text at the correct IEEE font size.")
    run.font.size = Pt(10)
    doc.add_paragraph("ABSTRACT A well-formed test abstract.")
    doc.add_paragraph("I. INTRODUCTION")
    doc.add_paragraph("REFERENCES")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    r = client.post(
        f"/api/submissions/{sub_id}/upload",
        files={"file": ("paper.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers=_auth(res_token),
    )
    assert r.status_code == 200

    r = client.get(f"/api/submissions/{sub_id}/ai-report", headers=_auth(res_token))
    reports = r.json()
    report_by_type = {rep["check_type"]: rep for rep in reports}
    assert "format" in report_by_type
    assert report_by_type["format"]["status"] == "complete"

    result = json.loads(report_by_type["format"]["result_json"])
    assert result["status"] == "complete"
    assert result["score"] == 100.0  # every measurable check on this well-formed doc should pass
    assert result["issues"] == []
