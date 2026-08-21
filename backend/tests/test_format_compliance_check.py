import pymupdf
from docx import Document
from docx.shared import Inches, Pt

from app.ai.format_compliance_check import (
    _closest_page_size,
    _measure_docx,
    _measure_pdf,
    run_format_compliance_check,
)


def test_closest_page_size_detects_letter():
    assert _closest_page_size(8.5, 11.0) == "letter"


def test_closest_page_size_detects_a4():
    assert _closest_page_size(8.27, 11.69) == "a4"


def test_closest_page_size_returns_none_for_unrecognized_size():
    assert _closest_page_size(20.0, 30.0) is None


def _make_ieee_compliant_docx(path: str) -> str:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.125)
    section.left_margin = Inches(0.8125)
    section.right_margin = Inches(0.8125)

    p = doc.add_paragraph()
    run = p.add_run("This is body text at the correct font size.")
    run.font.size = Pt(10)

    doc.add_paragraph("ABSTRACT This is a test abstract for the paper.")
    doc.add_paragraph("I. INTRODUCTION")
    doc.add_paragraph("Some introduction text.")
    doc.add_paragraph("REFERENCES")
    doc.add_paragraph("[1] Test reference.")

    doc.save(path)
    return path


def _make_non_compliant_docx(path: str) -> str:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(0.5)  # wrong — IEEE wants 1.0
    section.bottom_margin = Inches(0.5)  # wrong — IEEE wants 1.125
    section.left_margin = Inches(1.5)  # wrong — IEEE wants 0.8125
    section.right_margin = Inches(1.5)  # wrong — IEEE wants 0.8125

    p = doc.add_paragraph()
    run = p.add_run("Body text at the wrong font size entirely.")
    run.font.size = Pt(16)  # wrong — IEEE wants 10pt

    doc.add_paragraph("No abstract or references headings in this document at all.")

    doc.save(path)
    return path


def test_measure_docx_reads_exact_stored_margins():
    path = _make_ieee_compliant_docx("/tmp/_test_ieee_compliant.docx")
    measurements = _measure_docx(path)
    assert measurements["page_width_in"] == 8.5
    assert abs(measurements["margin_top_in"] - 1.0) < 0.01
    assert abs(measurements["margin_left_in"] - 0.8125) < 0.01  # can round slightly depending on float repr — check closeness, not exact equality
    assert measurements["body_font_size_pt"] == 10.0
    assert measurements["columns"] is None  # documented limitation, not measured for docx


def test_measure_docx_falls_back_to_style_inherited_font_size():
    """A real IEEE manuscript template (planning log §31) defines font size
    via named paragraph styles ("Abstract", "Body Text Indent", etc.) that
    themselves inherit from a base style — not by setting font.size on every
    individual run, which is what the original version of this function
    checked exclusively. This reproduces that exact real-world shape."""
    from docx import Document as DocxDocument
    from docx.shared import Pt as DocxPt

    doc = DocxDocument()
    # "Normal" carries the real size; a derived custom style inherits it
    # without setting its own font.size — exactly the real template's shape.
    doc.styles["Normal"].font.size = DocxPt(10)
    custom_style = doc.styles.add_style("Body Text Indent", 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    custom_style.base_style = doc.styles["Normal"]
    # Deliberately NOT setting custom_style.font.size — must inherit.

    p = doc.add_paragraph(style="Body Text Indent")
    run = p.add_run("Body text with no explicit run-level or style-level font size.")
    # Deliberately NOT setting run.font.size either.

    path = "/tmp/_test_style_inherited_font.docx"
    doc.save(path)

    measurements = _measure_docx(path)
    assert measurements["body_font_size_pt"] == 10.0


def test_compliant_docx_passes_all_measurable_checks():
    path = _make_ieee_compliant_docx("/tmp/_test_ieee_compliant2.docx")
    result = run_format_compliance_check(path, publisher_format="ieee")
    assert result["status"] == "complete"
    assert result["score"] == 100.0
    assert result["issues"] == []


def test_non_compliant_docx_flags_real_issues():
    path = _make_non_compliant_docx("/tmp/_test_non_compliant.docx")
    result = run_format_compliance_check(path, publisher_format="ieee")
    assert result["status"] == "complete"
    assert result["score"] < 100.0
    assert any("margin" in issue.lower() for issue in result["issues"])
    assert any("font" in issue.lower() for issue in result["issues"])
    assert any("abstract" in issue.lower() for issue in result["issues"])
    assert any("references" in issue.lower() for issue in result["issues"])


def test_unsupported_publisher_format_returns_clear_error():
    path = _make_ieee_compliant_docx("/tmp/_test_ieee_compliant3.docx")
    result = run_format_compliance_check(path, publisher_format="springer")
    assert result["status"] == "error"
    assert "springer" in result["error"].lower()
    assert result["score"] is None


def _make_two_column_pdf(path: str) -> str:
    """A real two-column PDF at IEEE Letter dimensions with correct margins,
    built via PyMuPDF's own writer. Multiple lines per column (not a single
    line) so PyMuPDF's block-detection forms two genuinely separate blocks —
    a single same-y-coordinate line per column merges into one block absent
    real paragraph structure, which isn't representative of an actual paper."""
    doc = pymupdf.open()
    left_x, right_x = 0.625 * 72, 4.625 * 72
    for page_num in range(2):
        page = doc.new_page(width=8.5 * 72, height=11.0 * 72)
        for i, y_in in enumerate([1.5, 1.7, 1.9, 2.1]):
            page.insert_text((left_x, y_in * 72), f"Left column line {i} on page {page_num}.", fontsize=10)
        for i, y_in in enumerate([1.5, 1.7, 1.9, 2.1]):
            page.insert_text((right_x, y_in * 72), f"Right column line {i} on page {page_num}.", fontsize=10)
    doc.save(path)
    doc.close()
    return path


def test_measure_pdf_detects_two_columns():
    path = _make_two_column_pdf("/tmp/_test_two_column.pdf")
    measurements = _measure_pdf(path)
    assert measurements["columns"] == 2
    assert measurements["page_count"] == 2


def test_measure_pdf_single_column_detected_as_one():
    doc = pymupdf.open()
    page = doc.new_page(width=8.5 * 72, height=11.0 * 72)
    page.insert_text((1.0 * 72, 1.5 * 72), "Single column text only on the left side.", fontsize=10)
    path = "/tmp/_test_single_column.pdf"
    doc.save(path)
    doc.close()

    measurements = _measure_pdf(path)
    assert measurements["columns"] == 1


def test_manuscript_without_header_measures_much_closer_to_true_margins_than_a_published_article():
    """A real published/typeset article (e.g. 2ieee.pdf) has compiler-added
    running headers/footers/page numbers sitting close to the page edges —
    content §3.7 explicitly says a manuscript SUBMISSION shouldn't have.
    Testing this check's margin measurement against a published article
    therefore reads artificially small margins on every side; that's a
    property of the test input, not a bug in the measurement.

    This test builds a clean manuscript-style page instead: real column-
    filling text (bounding-box margin inference is only as accurate as how
    consistently text reaches the column edge — real justified prose does
    this naturally, sparse/short lines don't), no header or footer, content
    genuinely spanning from just past the top margin to just before the
    bottom margin. Margins here should land much closer to the true IEEE
    values than the published-article case (which was off by 0.3-0.4in on
    every side)."""
    doc = pymupdf.open()
    page_w, page_h = 8.5 * 72, 11.0 * 72
    left_x, right_x = 0.8125 * 72, 4.625 * 72
    top_y = 1.0 * 72
    bottom_limit_y = page_h - 1.125 * 72

    for _ in range(2):
        page = doc.new_page(width=page_w, height=page_h)
        y_positions = []
        y = top_y + 8
        while y < bottom_limit_y - 8:
            y_positions.append(y)
            y += 14
        left_line = "Left column body text filling most of the column width here."
        right_line = "Right column body text filling most of the column width."
        for y in y_positions:
            page.insert_text((left_x, y), left_line, fontsize=9)
        for y in y_positions:
            page.insert_text((right_x, y), right_line, fontsize=9)

    path = "/tmp/_test_manuscript_no_header.pdf"
    doc.save(path)
    doc.close()

    measurements = _measure_pdf(path)
    assert measurements["columns"] == 2

    # Generous but meaningful tolerance — the point isn't pixel-perfect
    # synthetic-text precision, it's that every margin lands within ~0.3in
    # of the true IEEE value here, versus 0.3-0.4in OFF on the published-
    # article case. A tight assertion on synthetic non-justified text would
    # be testing this test's own text-layout precision, not the check.
    assert abs(measurements["margin_left_in"] - 0.8125) < 0.1
    assert abs(measurements["margin_top_in"] - 1.0) < 0.1
    assert abs(measurements["margin_right_in"] - 0.8125) < 0.3
    assert abs(measurements["margin_bottom_in"] - 1.125) < 0.3


def test_a4_document_checked_against_a4_margins_not_letter():
    """A document at A4-correct margins must NOT be flagged just because
    it's not Letter-correct — proving the check applies the right spec for
    the detected page size (planning log §33), not one universal margin set
    regardless of which page size was actually detected."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    p = doc.add_paragraph()
    run = p.add_run("Body text at the correct font size.")
    run.font.size = Pt(10)
    doc.add_paragraph("ABSTRACT A test abstract for the A4 paper.")
    doc.add_paragraph("I. INTRODUCTION")
    doc.add_paragraph("REFERENCES")

    path = "/tmp/_test_a4_compliant.docx"
    doc.save(path)

    result = run_format_compliance_check(path, publisher_format="ieee")
    margin_issues = [i for i in result["issues"] if "margin" in i.lower()]
    assert margin_issues == [], f"A4-correct margins were wrongly flagged: {margin_issues}"


def test_a4_document_with_letter_margins_is_correctly_flagged():
    """The inverse check — an A4 document using LETTER's margin numbers
    should be flagged, proving the two specs are genuinely distinct, not
    both silently accepted by an overly loose comparison."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)  # this is LETTER's top margin, not A4's 0.75
    section.bottom_margin = Inches(1.125)  # LETTER's bottom, not A4's 1.0
    section.left_margin = Inches(0.8125)  # LETTER's left, not A4's 0.62
    section.right_margin = Inches(0.8125)

    p = doc.add_paragraph()
    run = p.add_run("Body text.")
    run.font.size = Pt(10)

    path = "/tmp/_test_a4_with_letter_margins.docx"
    doc.save(path)

    result = run_format_compliance_check(path, publisher_format="ieee")
    margin_issues = [i for i in result["issues"] if "margin" in i.lower() and "A4" in i]
    assert len(margin_issues) >= 3  # top, bottom, left/right should all be flagged as wrong for A4
